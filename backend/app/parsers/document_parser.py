"""
DocumentParser
==============
Unified parser for .docx and .pdf files.
Returns a normalised dict per document:
  {filename, text, tables, headings, metadata}

Place at: backend/app/parsers/document_parser.py
Also create: backend/app/parsers/__init__.py  (empty)

Install deps:
  pip install python-docx pdfplumber pypdf
"""

import logging
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger("ai_buddy.parser")


class DocumentParser:

    async def parse(self, file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext == ".docx":
            return await self._parse_docx(path)
        elif ext == ".pdf":
            return await self._parse_pdf(path)
        else:
            raise ValueError(f"Unsupported file type for M1: {ext}. Use .docx or .pdf")

    # ── Word (.docx) ──────────────────────────────────────────────────────────

    async def _parse_docx(self, path: Path) -> Dict[str, Any]:
        import docx  # python-docx

        doc = docx.Document(str(path))
        text_parts: List[str] = []
        headings: List[Dict] = []
        tables: List[List[List[str]]] = []

        for para in doc.paragraphs:
            style = para.style.name
            if style.startswith("Heading"):
                level = int(style[-1]) if style[-1].isdigit() else 1
                headings.append({"level": level, "text": para.text.strip()})
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(rows)
                text_parts.append(self._table_to_text(rows))

        return {
            "filename": path.name,
            "text": "\n".join(text_parts),
            "headings": headings,
            "tables": tables,
            "metadata": {"source": "docx", "path": str(path)},
        }

    # ── PDF ───────────────────────────────────────────────────────────────────

    async def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        try:
            return await self._parse_pdf_pdfplumber(path)
        except ImportError:
            return await self._parse_pdf_pypdf(path)

    async def _parse_pdf_pdfplumber(self, path: Path) -> Dict[str, Any]:
        import pdfplumber

        text_parts: List[str] = []
        tables: List[List[List[str]]] = []

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
                for table in page.extract_tables():
                    if table:
                        tables.append(table)
                        text_parts.append(self._table_to_text(table))

        return {
            "filename": path.name,
            "text": "\n".join(text_parts),
            "headings": [],
            "tables": tables,
            "metadata": {"source": "pdf-pdfplumber", "path": str(path)},
        }

    async def _parse_pdf_pypdf(self, path: Path) -> Dict[str, Any]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text_parts = [p.extract_text() for p in reader.pages if p.extract_text()]

        return {
            "filename": path.name,
            "text": "\n".join(text_parts),
            "headings": [],
            "tables": [],
            "metadata": {"source": "pdf-pypdf", "path": str(path)},
        }

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _table_to_text(self, rows: List[List[str]]) -> str:
        return "\n".join(
            " | ".join(str(cell or "").strip() for cell in row)
            for row in rows
        )
