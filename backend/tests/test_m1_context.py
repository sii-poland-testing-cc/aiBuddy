"""
M1 Context Builder — unit + endpoint tests
===========================================
Run from backend/ with:
    pytest tests/test_m1_context.py -v
"""
from pathlib import Path

import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch


# ─────────────────────────────────────────────────────────────────────────────
# a) DocumentParser — .docx
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_document_parser_docx(tmp_path):
    """Parse an in-memory .docx and verify the returned structure."""
    import docx

    # Build a minimal Word document in memory
    doc = docx.Document()
    doc.add_heading("Test Process", level=1)
    doc.add_paragraph("This section describes the QA testing process.")
    doc.add_heading("Defect Management", level=2)
    doc.add_paragraph("Defects are logged and tracked by severity: Critical, High, Medium, Low.")

    docx_path = tmp_path / "test_doc.docx"
    doc.save(str(docx_path))

    from app.parsers.document_parser import DocumentParser
    result = await DocumentParser().parse(str(docx_path))

    assert result["filename"] == "test_doc.docx"
    assert len(result["text"]) > 0, "text should be non-empty"
    assert isinstance(result["headings"], list)
    assert len(result["headings"]) >= 2
    assert any("Test Process" in h["text"] for h in result["headings"])
    assert "metadata" in result


# ─────────────────────────────────────────────────────────────────────────────
# b) DocumentParser — .pdf
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_document_parser_pdf(sample_pdf_path):
    """Parse a minimal PDF and verify structure (text may be empty on simple PDFs)."""
    from app.parsers.document_parser import DocumentParser
    result = await DocumentParser().parse(sample_pdf_path)

    assert result["filename"].endswith(".pdf")
    # text can be empty for bare Type1 font PDFs — just assert structure is correct
    assert isinstance(result["text"], str)
    assert isinstance(result["headings"], list)
    assert isinstance(result["tables"], list)
    assert "metadata" in result
    assert result["metadata"]["source"] in ("pdf-pdfplumber", "pdf-pypdf")


# ─────────────────────────────────────────────────────────────────────────────
# c) ContextBuilder — index_from_docs / is_indexed
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_context_builder_index_from_docs():
    """index_from_docs returns chunk count > 0 and is_indexed becomes True."""
    from app.rag.context_builder import ContextBuilder

    cb = ContextBuilder()
    project_id = "test-unit-index"

    docs = [
        {
            "filename": "qa_process.docx",
            "text": (
                "Test Case: A set of conditions to verify system behaviour. "
                "Defect: A deviation from expected system behaviour. "
                "Coverage: The fraction of requirements exercised by tests. "
                "Regression Cycle: Periodic re-execution of tests after code changes. "
                "QA Engineer: Person responsible for quality assurance."
            ),
            "headings": [{"level": 1, "text": "QA Process"}],
            "tables": [],
            "metadata": {"source": "docx"},
        }
    ]

    count = await cb.index_from_docs(project_id, docs)
    assert count > 0, f"Expected chunk count > 0, got {count}"

    indexed = await cb.is_indexed(project_id)
    assert indexed is True


# ─────────────────────────────────────────────────────────────────────────────
# d) ContextBuilderWorkflow — mock LLM + parser
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_context_builder_workflow_mock():
    """
    Run ContextBuilderWorkflow with llm=None (mock data path).
    Patch DocumentParser.parse and ContextBuilder to avoid real I/O.
    """
    fixed_doc = {
        "filename": "domain.docx",
        "text": "QA process with test cases, defects, and coverage metrics.",
        "headings": [{"level": 1, "text": "QA Process"}],
        "tables": [],
        "metadata": {"source": "docx"},
    }

    mock_cb_instance = MagicMock()
    mock_cb_instance.index_from_docs = AsyncMock(return_value=4)

    mock_parser_instance = MagicMock()
    mock_parser_instance.parse = AsyncMock(return_value=fixed_doc)

    with (
        patch("app.agents.context_builder_workflow.ContextBuilder", return_value=mock_cb_instance),
        patch("app.agents.context_builder_workflow.DocumentParser", return_value=mock_parser_instance),
    ):
        from app.agents.context_builder_workflow import ContextBuilderWorkflow

        workflow = ContextBuilderWorkflow(llm=None, timeout=60)
        handler = workflow.run(
            project_id="test-unit-workflow",
            file_paths=["/fake/domain.docx"],
        )

        # Consume progress events, then collect result
        async for _ in handler.stream_events():
            pass
        result = await handler

    # Structural assertions
    assert result["rag_ready"] is True

    mm = result["mind_map"]
    assert "nodes" in mm and isinstance(mm["nodes"], list)
    assert "edges" in mm and isinstance(mm["edges"], list)
    assert len(mm["nodes"]) > 0
    assert len(mm["edges"]) > 0

    glossary = result["glossary"]
    assert isinstance(glossary, list)
    assert len(glossary) > 0
    assert all("term" in t and "definition" in t for t in glossary)

    stats = result["stats"]
    assert stats["entity_count"] == len(mm["nodes"])
    assert stats["term_count"] == len(glossary)


# ─────────────────────────────────────────────────────────────────────────────
# e) GET /api/context/{project_id}/status — nonexistent project
# ─────────────────────────────────────────────────────────────────────────────

def test_context_status_endpoint_nonexistent(app_client):
    """Status for an unknown project returns 200 with rag_ready=false."""
    r = app_client.get("/api/context/nonexistent-project-xyz/status")
    assert r.status_code == 200
    body = r.json()
    assert body["rag_ready"] is False
    assert body["artefacts_ready"] is False
    assert body["stats"] is None


# ─────────────────────────────────────────────────────────────────────────────
# f) POST /api/context/{project_id}/build — invalid file type
# ─────────────────────────────────────────────────────────────────────────────

def test_context_build_endpoint_invalid_file(app_client, tmp_path):
    """Uploading a .txt file to /build should return HTTP 400."""
    txt_file = tmp_path / "notes.txt"
    txt_file.write_text("This is not a Word document.")

    with txt_file.open("rb") as fh:
        r = app_client.post(
            "/api/context/test-project-invalid/build",
            files={"files": ("notes.txt", fh, "text/plain")},
        )

    assert r.status_code == 400
    detail = r.json().get("detail", "")
    assert "docx" in detail.lower() or "pdf" in detail.lower()


# ─────────────────────────────────────────────────────────────────────────────
# g) Artefacts persisted to DB and readable after cache clear
# ─────────────────────────────────────────────────────────────────────────────

def test_artefacts_persisted_to_db(app_client):
    """
    After a successful M1 build the artefacts must survive an in-memory
    cache flush — i.e. they are read back from the DB, not from _context_store.
    """
    from pathlib import Path
    import app.api.routes.context as context_module

    fixture = Path(__file__).parent / "fixtures" / "sample_domain.docx"
    assert fixture.exists(), f"Fixture missing: {fixture}"

    # 1. Create a project so the DB row exists for the FK-free write
    proj_r = app_client.post("/api/projects/", json={"name": "persist-test-project"})
    assert proj_r.status_code in (200, 201)
    project_id = proj_r.json()["project_id"]

    # 2. Run the full M1 build with a mocked LLM (no real API calls)
    with patch("app.api.routes.context.get_llm", return_value=_make_context_mock_llm()), \
         fixture.open("rb") as fh:
        build_r = app_client.post(
            f"/api/context/{project_id}/build",
            files={"files": (fixture.name, fh,
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert build_r.status_code == 200, f"Build failed: {build_r.text[:200]}"

    # Confirm artefacts landed in the cache
    assert project_id in context_module._context_store

    # 3. Evict from in-memory cache to force a DB read
    context_module._context_store.pop(project_id)
    assert project_id not in context_module._context_store

    # 4. GET /mindmap — must come from DB now
    mm_r = app_client.get(f"/api/context/{project_id}/mindmap")
    assert mm_r.status_code == 200, f"Mindmap after cache clear: {mm_r.text}"
    mm = mm_r.json()
    assert isinstance(mm["nodes"], list) and len(mm["nodes"]) > 0
    assert isinstance(mm["edges"], list) and len(mm["edges"]) > 0

    # 5. GET /glossary — also from DB
    context_module._context_store.pop(project_id, None)
    gl_r = app_client.get(f"/api/context/{project_id}/glossary")
    assert gl_r.status_code == 200
    glossary = gl_r.json()
    assert isinstance(glossary, list) and len(glossary) > 0

    # 6. GET /status — context_built_at should be present
    st_r = app_client.get(f"/api/context/{project_id}/status")
    assert st_r.status_code == 200
    status = st_r.json()
    assert status["artefacts_ready"] is True
    assert status["context_built_at"] is not None


# ─────────────────────────────────────────────────────────────────────────────
# h) Append mode merges artefacts from two builds
# ─────────────────────────────────────────────────────────────────────────────

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_SYNTHETIC = Path(__file__).parent / "fixtures" / "synthetic_docs"

import json as _json

_MOCK_ENTITIES = _json.dumps({
    "entities": [
        {"id": "e1", "name": "Test Case", "type": "data", "description": "A test scenario"},
        {"id": "e2", "name": "Defect",    "type": "data", "description": "A software defect"},
    ],
    "relations": [{"source": "e1", "target": "e2", "label": "reveals"}],
})
_MOCK_GLOSSARY = _json.dumps([
    {"term": "Test Case", "definition": "Conditions to verify behaviour.", "related_terms": [], "source": "docs"},
    {"term": "Defect",    "definition": "Deviation from expected behaviour.", "related_terms": [], "source": "docs"},
])
_MOCK_APPROVED = _json.dumps({"verdict": "APPROVED"})


def _make_context_mock_llm():
    """Return a mock LLM that returns deterministic responses for M1 extraction calls."""
    mock = MagicMock()

    async def _side(prompt, **kwargs):
        if "entities and their relationships" in prompt:
            return _MOCK_ENTITIES
        if "domain-specific term" in prompt:  # _enumerate_term_names (phase 1)
            return _json.dumps(["Test Case", "Defect", "QA Engineer", "Test Suite", "Coverage"])
        if "Write glossary definitions" in prompt:  # _define_term_group (phase 2)
            return _MOCK_GLOSSARY
        return _MOCK_APPROVED  # review / refine

    mock.acomplete = AsyncMock(side_effect=_side)
    return mock


def _build(client, project_id: str, filename: str, mode: str = "append") -> None:
    """POST a single .docx to /build and consume the SSE response."""
    path = _SYNTHETIC / filename
    assert path.exists(), f"Fixture missing: {path}"
    url = f"/api/context/{project_id}/build"
    if mode != "append":
        url += f"?mode={mode}"
    with patch("app.api.routes.context.get_llm", return_value=_make_context_mock_llm()), \
         path.open("rb") as fh:
        r = client.post(url, files={"files": (filename, fh, _DOCX_MIME)})
    assert r.status_code == 200, f"Build failed ({mode}): {r.text[:300]}"


def _glossary(client, project_id: str) -> list:
    r = client.get(f"/api/context/{project_id}/glossary")
    assert r.status_code == 200
    return r.json()


def test_append_mode_merges_artefacts(app_client):
    """
    Appending a second document merges its glossary with the first:
      - no duplicate terms
      - merged term count > either file built alone

    Build 1 uses the standard mock (5 terms).  Build 2 uses an extended mock
    that adds 2 unique terms absent from build 1, so the merged glossary is 7
    terms — strictly more than either single build.
    """
    def _make_extended_mock_llm():
        """Standard mock plus 2 extra terms unique to the second document."""
        mock = MagicMock()
        extended_terms = ["Test Case", "Defect", "QA Engineer", "Test Suite", "Coverage",
                          "Payment Gateway", "Transaction Ledger"]
        extended_glossary = _json.dumps([
            {"term": "Test Case",          "definition": "Conditions to verify behaviour.", "related_terms": [], "source": "docs"},
            {"term": "Defect",             "definition": "Deviation from expected behaviour.", "related_terms": [], "source": "docs"},
            {"term": "QA Engineer",        "definition": "Quality assurance specialist.", "related_terms": [], "source": "docs"},
            {"term": "Test Suite",         "definition": "Collection of test cases.", "related_terms": [], "source": "docs"},
            {"term": "Coverage",           "definition": "Fraction of requirements exercised.", "related_terms": [], "source": "docs"},
            {"term": "Payment Gateway",    "definition": "External payment processor.", "related_terms": [], "source": "docs"},
            {"term": "Transaction Ledger", "definition": "Ledger of all payment transactions.", "related_terms": [], "source": "docs"},
        ])

        async def _side(prompt, **kwargs):
            if "entities and their relationships" in prompt:
                return _MOCK_ENTITIES
            if "domain-specific term" in prompt:
                return _json.dumps(extended_terms)
            if "Write glossary definitions" in prompt:
                return extended_glossary
            return _MOCK_APPROVED

        mock.acomplete = AsyncMock(side_effect=_side)
        return mock

    r = app_client.post("/api/projects/", json={"name": "append-merge-test"})
    assert r.status_code in (200, 201)
    pid = r.json()["project_id"]

    # Build 1: srs_payment_module.docx — standard mock (5 terms)
    _build(app_client, pid, "srs_payment_module.docx")
    n_srs = len(_glossary(app_client, pid))

    # Build 2: qa_process.docx (append) — extended mock adds 2 unique extra terms
    path = _SYNTHETIC / "qa_process.docx"
    assert path.exists()
    with patch("app.api.routes.context.get_llm", return_value=_make_extended_mock_llm()), \
         path.open("rb") as fh:
        r2 = app_client.post(
            f"/api/context/{pid}/build",
            files={"files": ("qa_process.docx", fh, _DOCX_MIME)},
        )
    assert r2.status_code == 200, f"Build 2 failed: {r2.text[:300]}"
    merged = _glossary(app_client, pid)

    # No duplicate terms (case-insensitive)
    lowered = [t["term"].lower() for t in merged]
    assert len(set(lowered)) == len(merged), "Merged glossary contains duplicate terms"

    # Merged count strictly greater than either file alone
    assert len(merged) > n_srs, f"Expected merged ({len(merged)}) > srs alone ({n_srs})"


# ─────────────────────────────────────────────────────────────────────────────
# i) Rebuild mode replaces artefacts — no trace of the previous build
# ─────────────────────────────────────────────────────────────────────────────

def test_rebuild_mode_replaces_artefacts(app_client):
    """
    After a rebuild the glossary reflects only the new document,
    and document_count resets to 1.
    """
    r = app_client.post("/api/projects/", json={"name": "rebuild-test"})
    assert r.status_code in (200, 201)
    pid = r.json()["project_id"]

    # Initial build with srs_payment_module.docx
    _build(app_client, pid, "srs_payment_module.docx")

    # Rebuild with qa_process.docx only
    _build(app_client, pid, "qa_process.docx", mode="rebuild")
    glossary = _glossary(app_client, pid)
    terms_after = {t["term"].lower() for t in glossary}

    # At least one QA-process term present (mock data always has "defect")
    qa_terms = {"defect", "severity", "bug", "defect lifecycle", "qa engineer"}
    assert terms_after & qa_terms, (
        f"Expected at least one QA-process term in {terms_after}"
    )

    # document_count reset to 1 and only qa_process.docx listed
    status = app_client.get(f"/api/context/{pid}/status").json()
    assert status["document_count"] == 1, (
        f"Expected document_count=1 after rebuild, got {status['document_count']}"
    )
    cf_names = [f["name"] for f in status["context_files"]]
    assert cf_names == ["qa_process.docx"], (
        f"Expected only qa_process.docx, got {cf_names}"
    )


# ─────────────────────────────────────────────────────────────────────────────
# j) context_files are tracked across multiple builds
# ─────────────────────────────────────────────────────────────────────────────

def test_context_files_tracked(app_client):
    """
    /status returns document_count and context_files correctly after each build.
    """
    r = app_client.post("/api/projects/", json={"name": "files-tracking-test"})
    assert r.status_code in (200, 201)
    pid = r.json()["project_id"]

    # Build 1
    _build(app_client, pid, "srs_payment_module.docx")
    s1 = app_client.get(f"/api/context/{pid}/status").json()
    assert s1["document_count"] == 1
    srs_entry = next(f for f in s1["context_files"] if f["name"] == "srs_payment_module.docx")
    assert srs_entry["indexed_at"] is not None, "indexed_at must be set after build"

    # Build 2 (append) — srs already indexed, only test_plan should be processed
    _build(app_client, pid, "test_plan_payment.docx", mode="append")
    s2 = app_client.get(f"/api/context/{pid}/status").json()
    assert s2["document_count"] == 2
    names2 = {f["name"] for f in s2["context_files"]}
    assert "srs_payment_module.docx" in names2
    assert "test_plan_payment.docx" in names2
    # Both entries must carry indexed_at
    for f in s2["context_files"]:
        assert f["indexed_at"] is not None, f"{f['name']} missing indexed_at"


# ─────────────────────────────────────────────────────────────────────────────
# AuditSnapshot model smoke test
# ─────────────────────────────────────────────────────────────────────────────

def test_audit_snapshot_table_exists(app_client):
    """
    app_client triggers lifespan → init_db() which must create audit_snapshots.
    Verify the model class and that all JSON fields round-trip correctly.
    """
    from app.db.models import AuditSnapshot

    assert AuditSnapshot.__tablename__ == "audit_snapshots"

    summary_data = {
        "coverage_pct": 33.3,
        "duplicates_found": 1,
        "requirements_total": 12,
        "requirements_covered": 4,
        "untagged_cases": 2,
    }
    uncovered = ["FR-001", "FR-005", "FR-007", "FR-008", "FR-009", "FR-010", "FR-011", "FR-012"]
    recs = ["Add tests for FR-001", "Add tests for FR-005"]
    diff_data = {
        "coverage_delta": 12.5,
        "duplicates_delta": -2,
        "new_covered": ["FR-003"],
        "newly_uncovered": [],
        "files_added": ["v17.xlsx"],
        "files_removed": ["v16.xlsx"],
    }

    import uuid as _uuid
    snap_id = str(_uuid.uuid4())
    snap = AuditSnapshot(
        id=snap_id,
        project_id="test-project-id",
        files_used=["file1.xlsx"],
        summary=summary_data,
        requirements_uncovered=uncovered,
        recommendations=recs,
        diff=diff_data,
    )

    # All JSON fields are stored as Python objects (JsonType handles serialization at DB boundary)
    assert snap.summary["coverage_pct"] == 33.3             # type: ignore[index]
    assert snap.requirements_uncovered == uncovered          # type: ignore[comparison-overlap]
    assert snap.recommendations == recs                      # type: ignore[comparison-overlap]
    assert snap.diff["coverage_delta"] == 12.5               # type: ignore[index]
    assert snap.files_used == ["file1.xlsx"]                 # type: ignore[comparison-overlap]

    # id matches what we passed
    assert snap.id == snap_id
    assert len(snap.id) == 36   # UUID4 string


# ─────────────────────────────────────────────────────────────────────────────
# audit_workflow_integration — coverage helpers
# (previously tested via AuditWorkflow._requirements_in_tests / _extract_requirements,
#  now tested at their live home in the integration module)
# ─────────────────────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_requirements_in_tests_pattern_match():
    """Pattern matching covers requirements mentioned in test case fields."""
    from app.agents.audit_workflow_integration import _match_requirements_to_tests

    cases = [{"name": "Test FR-002 card schemes", "tags": ""}]
    known_reqs = ["FR-001", "FR-002", "FR-003"]

    result = await _match_requirements_to_tests(cases, known_reqs, [], llm=None)

    assert "FR-002" in result
    assert "FR-001" not in result
    assert "FR-003" not in result


@pytest.mark.asyncio
async def test_extract_requirements_returns_list():
    """
    _legacy_extract with llm=None returns the mock list.
    With a real-ish mock LLM, it parses JSON and returns a list of strings.
    """
    from unittest.mock import AsyncMock, MagicMock
    from app.agents.audit_workflow_integration import _legacy_extract

    rag_context = (
        "FR-001: The system shall support Visa payments.\n"
        "FR-002: The system shall support Mastercard payments.\n"
        "FR-003: The capture window is 7 days by default.\n"
    )

    # 1. llm=None → empty list (no phantom IDs)
    result = await _legacy_extract(rag_context, llm=None)
    assert isinstance(result, list), "Expected a list"
    assert result == [], "Expected empty list when no LLM available"

    # 2. With a mock LLM that returns a JSON array
    mock_llm = MagicMock()
    mock_llm.acomplete = AsyncMock(return_value='["FR-001", "FR-002", "FR-003"]')
    result2 = await _legacy_extract(rag_context, llm=mock_llm)
    assert isinstance(result2, list)
    assert all(isinstance(r, str) for r in result2)
    assert "FR-001" in result2
    assert "FR-003" in result2


# ─────────────────────────────────────────────────────────────────────────────
# _merge_mind_maps + _merge_glossaries  (pure-function unit tests)
# ─────────────────────────────────────────────────────────────────────────────

def test_merge_mind_maps_deduplicates_nodes_by_id():
    from app.api.routes.context import _merge_mind_maps

    existing = {"nodes": [{"id": "A", "label": "A"}], "edges": []}
    new      = {"nodes": [{"id": "A", "label": "A-new"}, {"id": "B", "label": "B"}], "edges": []}
    result = _merge_mind_maps(existing, new)

    ids = [n["id"] for n in result["nodes"]]
    assert ids.count("A") == 1, "Duplicate node A must appear exactly once"
    assert "B" in ids
    # existing wins on duplicate (existing processed first)
    assert next(n for n in result["nodes"] if n["id"] == "A")["label"] == "A"


def test_merge_mind_maps_deduplicates_edges_by_source_target():
    from app.api.routes.context import _merge_mind_maps

    e1 = {"source": "A", "target": "B"}
    existing = {"nodes": [], "edges": [e1]}
    new      = {"nodes": [], "edges": [{"source": "A", "target": "B"}, {"source": "B", "target": "C"}]}
    result = _merge_mind_maps(existing, new)

    keys = [(e["source"], e["target"]) for e in result["edges"]]
    assert keys.count(("A", "B")) == 1
    assert ("B", "C") in keys


def test_merge_mind_maps_empty_inputs():
    from app.api.routes.context import _merge_mind_maps

    assert _merge_mind_maps({}, {}) == {"nodes": [], "edges": []}
    assert _merge_mind_maps({"nodes": [{"id": "X"}], "edges": []}, {}) == {
        "nodes": [{"id": "X"}], "edges": []
    }


def test_merge_glossaries_new_wins_on_duplicate_term():
    from app.api.routes.context import _merge_glossaries

    existing = [{"term": "Auth", "definition": "old"}]
    new      = [{"term": "auth", "definition": "new"}, {"term": "Token", "definition": "JWT"}]
    result = _merge_glossaries(existing, new)

    terms = {t["term"].lower(): t for t in result}
    assert terms["auth"]["definition"] == "new", "New entry must win on duplicate term"
    assert "token" in terms


def test_merge_glossaries_case_insensitive_dedup():
    from app.api.routes.context import _merge_glossaries

    existing = [{"term": "SMOKE TEST", "definition": "a"}]
    new      = [{"term": "smoke test", "definition": "b"}]
    result = _merge_glossaries(existing, new)

    assert len(result) == 1, "Case-insensitive duplicate should produce one entry"
    assert result[0]["definition"] == "b"


def test_merge_glossaries_empty_inputs():
    from app.api.routes.context import _merge_glossaries

    assert _merge_glossaries([], []) == []
    assert _merge_glossaries([{"term": "T", "definition": "d"}], []) == [
        {"term": "T", "definition": "d"}
    ]


# ─────────────────────────────────────────────────────────────────────────────
# k) _parse_context_files — all legacy and current formats
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_context_files_none():
    from app.api.routes.context import _parse_context_files
    assert _parse_context_files(None) == {"docs": [], "jira": []}


def test_parse_context_files_unknown_type():
    from app.api.routes.context import _parse_context_files
    assert _parse_context_files(42) == {"docs": [], "jira": []}


def test_parse_context_files_legacy_list():
    """Oldest format: plain strings, jira items have 'jira:' prefix."""
    from app.api.routes.context import _parse_context_files
    raw = ["doc.pdf", "jira:PROJ-1", "doc2.docx"]
    result = _parse_context_files(raw)
    assert len(result["docs"]) == 2
    assert result["docs"][0] == {"name": "doc.pdf", "indexed_at": None}
    assert result["docs"][1] == {"name": "doc2.docx", "indexed_at": None}
    assert len(result["jira"]) == 1
    assert result["jira"][0]["key"] == "PROJ-1"
    assert result["jira"][0]["indexed_at"] is None


def test_parse_context_files_dict_v1_string_docs():
    """Dict format with docs as plain strings (no indexed_at yet)."""
    from app.api.routes.context import _parse_context_files
    raw = {"docs": ["a.docx", "b.pdf"], "jira": []}
    result = _parse_context_files(raw)
    assert result["docs"] == [
        {"name": "a.docx", "indexed_at": None},
        {"name": "b.pdf",  "indexed_at": None},
    ]


def test_parse_context_files_dict_v2_object_docs():
    """Current format: docs already have {name, indexed_at} shape — pass through."""
    from app.api.routes.context import _parse_context_files
    raw = {"docs": [{"name": "a.docx", "indexed_at": "2025-01-01T00:00:00+00:00"}], "jira": []}
    result = _parse_context_files(raw)
    assert result["docs"] == [{"name": "a.docx", "indexed_at": "2025-01-01T00:00:00+00:00"}]


def test_parse_context_files_mixed_docs():
    """Dict with some string docs and some object docs — both normalised."""
    from app.api.routes.context import _parse_context_files
    raw = {"docs": ["plain.docx", {"name": "obj.pdf", "indexed_at": "2025-06-01T00:00:00Z"}], "jira": []}
    result = _parse_context_files(raw)
    assert result["docs"][0] == {"name": "plain.docx", "indexed_at": None}
    assert result["docs"][1] == {"name": "obj.pdf", "indexed_at": "2025-06-01T00:00:00Z"}


# ─────────────────────────────────────────────────────────────────────────────
# l) _jira_mds_for_append — filters by indexed_at vs context_built_at
# ─────────────────────────────────────────────────────────────────────────────

def test_jira_mds_for_append_no_jira_dir(tmp_path):
    """No jira/ directory → empty result."""
    import app.api.routes.context as ctx_module
    from unittest.mock import patch

    with patch.object(ctx_module, "UPLOAD_ROOT", tmp_path):
        result = ctx_module._jira_mds_for_append("pid", {"jira": []}, None)
    assert result == []


def test_jira_mds_for_append_no_prior_build(tmp_path):
    """context_built_at=None → all items with an existing .md are included."""
    import app.api.routes.context as ctx_module
    from unittest.mock import patch

    jira_dir = tmp_path / "pid" / "context" / "jira"
    jira_dir.mkdir(parents=True)
    (jira_dir / "PROJ-1.md").write_text("# PROJ-1")

    cf = {"jira": [{"key": "PROJ-1", "indexed_at": "2025-01-01T00:00:00+00:00"}]}
    with patch.object(ctx_module, "UPLOAD_ROOT", tmp_path):
        result = ctx_module._jira_mds_for_append("pid", cf, context_built_at=None)
    assert len(result) == 1


def test_jira_mds_for_append_filters_stale_items(tmp_path):
    """Items indexed before context_built_at are excluded; newer ones included."""
    import app.api.routes.context as ctx_module
    from datetime import datetime, timezone
    from unittest.mock import patch

    jira_dir = tmp_path / "pid" / "context" / "jira"
    jira_dir.mkdir(parents=True)
    (jira_dir / "PROJ-1.md").write_text("# PROJ-1")  # indexed after → include
    (jira_dir / "PROJ-2.md").write_text("# PROJ-2")  # indexed before → exclude

    context_built_at = datetime(2025, 6, 1, 12, 0, 0, tzinfo=timezone.utc)
    cf = {"jira": [
        {"key": "PROJ-1", "indexed_at": "2025-06-15T00:00:00+00:00"},  # after
        {"key": "PROJ-2", "indexed_at": "2025-05-01T00:00:00+00:00"},  # before
    ]}
    with patch.object(ctx_module, "UPLOAD_ROOT", tmp_path):
        result = ctx_module._jira_mds_for_append("pid", cf, context_built_at)

    assert len(result) == 1
    assert "PROJ-1.md" in result[0]


def test_jira_mds_for_append_missing_md_skipped(tmp_path):
    """Items without a .md file on disk are skipped silently."""
    import app.api.routes.context as ctx_module
    from unittest.mock import patch

    (tmp_path / "pid" / "context" / "jira").mkdir(parents=True)
    # PROJ-1.md does NOT exist on disk

    cf = {"jira": [{"key": "PROJ-1", "indexed_at": None}]}
    with patch.object(ctx_module, "UPLOAD_ROOT", tmp_path):
        result = ctx_module._jira_mds_for_append("pid", cf, context_built_at=None)
    assert result == []


# ─────────────────────────────────────────────────────────────────────────────
# m) Append noop — re-uploading an already-indexed file yields type:"noop"
# ─────────────────────────────────────────────────────────────────────────────

def _parse_sse_events(response_text: str) -> list:
    """Extract parsed JSON objects from an SSE response body."""
    import json
    events = []
    for line in response_text.splitlines():
        if line.startswith("data: ") and not line.startswith("data: [DONE]"):
            try:
                events.append(json.loads(line[6:]))
            except json.JSONDecodeError:
                pass
    return events


def test_append_noop_when_file_already_indexed(app_client):
    """
    Second append build with the same already-indexed file must return a noop
    SSE event and leave the context unchanged.
    """
    r = app_client.post("/api/projects/", json={"name": "noop-test"})
    pid = r.json()["project_id"]

    # Build 1 — file gets indexed
    _build(app_client, pid, "srs_payment_module.docx")
    status_before = app_client.get(f"/api/context/{pid}/status").json()

    # Build 2 — same file, append mode → should trigger noop
    path = _SYNTHETIC / "srs_payment_module.docx"
    with patch("app.api.routes.context.get_llm", return_value=_make_context_mock_llm()), \
         path.open("rb") as fh:
        r2 = app_client.post(
            f"/api/context/{pid}/build",
            files={"files": ("srs_payment_module.docx", fh, _DOCX_MIME)},
        )
    assert r2.status_code == 200

    events = _parse_sse_events(r2.text)
    event_types = [e.get("type") for e in events]
    assert "noop" in event_types, f"Expected noop event; got event types: {event_types}"

    noop_ev = next(e for e in events if e.get("type") == "noop")
    assert noop_ev["data"]["message"], "noop event must carry a non-empty message"

    # Status must be unchanged (no extra build happened)
    status_after = app_client.get(f"/api/context/{pid}/status").json()
    assert status_after["document_count"] == status_before["document_count"]


# ─────────────────────────────────────────────────────────────────────────────
# n) indexed_at < context_built_at invariant holds after every build
# ─────────────────────────────────────────────────────────────────────────────

def test_indexed_at_earlier_than_context_built_at(app_client):
    """
    After a build, every doc entry must have indexed_at strictly earlier than
    context_built_at — this ensures the append-skip filter works correctly.
    """
    from datetime import datetime

    r = app_client.post("/api/projects/", json={"name": "invariant-test"})
    pid = r.json()["project_id"]
    _build(app_client, pid, "srs_payment_module.docx")

    status = app_client.get(f"/api/context/{pid}/status").json()
    assert status["context_built_at"] is not None

    cba = datetime.fromisoformat(status["context_built_at"]).replace(tzinfo=None)

    for f in status["context_files"]:
        assert f["indexed_at"] is not None, f"{f['name']} has no indexed_at"
        indexed_dt = datetime.fromisoformat(f["indexed_at"]).replace(tzinfo=None)
        assert indexed_dt < cba, (
            f"{f['name']}: indexed_at ({indexed_dt}) must be < context_built_at ({cba})"
        )


# ─────────────────────────────────────────────────────────────────────────────
# o) Jira context endpoints — add / delete (JiraClient mocked)
# ─────────────────────────────────────────────────────────────────────────────

_MOCK_JIRA_DATA = {
    "root": {
        "key": "PROJ-1", "summary": "Test story", "description": "Acceptance criteria here",
        "type": "Story", "status": "In Progress", "priority": "High", "labels": [],
    },
    "parent": None,
    "children": [],
    "linked": [],
}


def test_add_context_jira_missing_config(app_client):
    """Returns 400 when the project has no Jira credentials configured."""
    r = app_client.post("/api/projects/", json={"name": "jira-no-config"})
    pid = r.json()["project_id"]

    resp = app_client.post(f"/api/context/{pid}/jira", json={"issue_key": "PROJ-1"})
    assert resp.status_code == 400


def test_add_context_jira_not_found(app_client):
    """Returns 404 when JiraClient returns None (issue not found)."""
    r = app_client.post("/api/projects/", json={"name": "jira-404-test"})
    pid = r.json()["project_id"]
    app_client.put(f"/api/projects/{pid}/settings", json={
        "jira_url": "https://fake.atlassian.net",
        "jira_user_email": "test@test.com",
        "jira_api_key": "fake-key",
    })

    with patch("app.api.routes.context.JiraClient") as MockJira:
        MockJira.return_value.fetch_with_context = AsyncMock(return_value=None)
        resp = app_client.post(f"/api/context/{pid}/jira", json={"issue_key": "PROJ-99"})

    assert resp.status_code == 404


def test_add_context_jira_success(app_client):
    """With mocked client, issue is added to context_files and appears in status."""
    r = app_client.post("/api/projects/", json={"name": "jira-add-test"})
    pid = r.json()["project_id"]
    app_client.put(f"/api/projects/{pid}/settings", json={
        "jira_url": "https://fake.atlassian.net",
        "jira_user_email": "test@test.com",
        "jira_api_key": "fake-key",
    })

    with patch("app.api.routes.context.JiraClient") as MockJira, \
         patch("app.api.routes.context._context_builder") as mock_cb:
        MockJira.return_value.fetch_with_context = AsyncMock(return_value=_MOCK_JIRA_DATA)
        mock_cb.index_files = AsyncMock()

        resp = app_client.post(f"/api/context/{pid}/jira", json={"issue_key": "proj-1"})

    assert resp.status_code == 201
    body = resp.json()
    assert body["issue_key"] == "PROJ-1"  # uppercased
    assert body["indexed"] is True
    assert body["indexed_at"] is not None

    status = app_client.get(f"/api/context/{pid}/status").json()
    assert any(j["key"] == "PROJ-1" for j in status["jira_sources"])


def test_add_context_jira_idempotent(app_client):
    """Adding the same issue twice updates indexed_at rather than duplicating."""
    r = app_client.post("/api/projects/", json={"name": "jira-idempotent-test"})
    pid = r.json()["project_id"]
    app_client.put(f"/api/projects/{pid}/settings", json={
        "jira_url": "https://fake.atlassian.net",
        "jira_user_email": "u@u.com",
        "jira_api_key": "k",
    })

    for _ in range(2):
        with patch("app.api.routes.context.JiraClient") as MockJira, \
             patch("app.api.routes.context._context_builder") as mock_cb:
            MockJira.return_value.fetch_with_context = AsyncMock(return_value=_MOCK_JIRA_DATA)
            mock_cb.index_files = AsyncMock()
            app_client.post(f"/api/context/{pid}/jira", json={"issue_key": "PROJ-1"})

    status = app_client.get(f"/api/context/{pid}/status").json()
    proj1_entries = [j for j in status["jira_sources"] if j["key"] == "PROJ-1"]
    assert len(proj1_entries) == 1, "Duplicate add must not create two entries"


def test_delete_context_jira_success(app_client):
    """Deleting an existing Jira source removes it from context_files."""
    r = app_client.post("/api/projects/", json={"name": "jira-delete-test"})
    pid = r.json()["project_id"]
    app_client.put(f"/api/projects/{pid}/settings", json={
        "jira_url": "https://fake.atlassian.net",
        "jira_user_email": "u@u.com",
        "jira_api_key": "k",
    })

    with patch("app.api.routes.context.JiraClient") as MockJira, \
         patch("app.api.routes.context._context_builder") as mock_cb:
        MockJira.return_value.fetch_with_context = AsyncMock(return_value=_MOCK_JIRA_DATA)
        mock_cb.index_files = AsyncMock()
        app_client.post(f"/api/context/{pid}/jira", json={"issue_key": "PROJ-1"})

    del_resp = app_client.delete(f"/api/context/{pid}/jira/PROJ-1")
    assert del_resp.status_code == 204

    status = app_client.get(f"/api/context/{pid}/status").json()
    assert not any(j["key"] == "PROJ-1" for j in status["jira_sources"])


def test_delete_context_jira_not_found(app_client):
    """Deleting a non-existent Jira key returns 404."""
    r = app_client.post("/api/projects/", json={"name": "jira-del-404"})
    pid = r.json()["project_id"]
    resp = app_client.delete(f"/api/context/{pid}/jira/PROJ-99")
    assert resp.status_code == 404


# ─────────────────────────────────────────────────────────────────────────────
# p) /status always returns jira_sources field
# ─────────────────────────────────────────────────────────────────────────────

def test_status_includes_jira_sources_field(app_client):
    """jira_sources is always present in /status — empty list for a fresh project."""
    r = app_client.post("/api/projects/", json={"name": "jira-sources-field-test"})
    pid = r.json()["project_id"]

    status = app_client.get(f"/api/context/{pid}/status").json()
    assert "jira_sources" in status
    assert isinstance(status["jira_sources"], list)
