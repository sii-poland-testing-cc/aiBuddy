"""
Generate synthetic QA documentation .docx fixtures for M1 Context Builder testing.
Run from the repo root or from backend/:
    python backend/tests/fixtures/generate_synthetic_docs.py
"""

import asyncio
import sys
from pathlib import Path

import docx

OUT_DIR = Path(__file__).parent / "synthetic_docs"
OUT_DIR.mkdir(exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _heading(doc, text: str, level: int):
    doc.add_heading(text, level=level)


def _para(doc, text: str):
    doc.add_paragraph(text)


def _table(doc, headers: list[str], rows: list[list[str]]):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        tbl_row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            tbl_row.cells[c_idx].text = val
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# 1. srs_payment_module.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_srs(path: Path):
    doc = docx.Document()
    doc.core_properties.title = "PayFlow SRS"
    doc.core_properties.author = "QA Team"

    # Title
    title = doc.add_heading("Software Requirements Specification — PayFlow Payment Module", 0)

    # ── System Overview ───────────────────────────────────────────────────────
    _heading(doc, "System Overview", 1)
    _para(doc,
        "PayFlow is an enterprise-grade payment processing gateway designed to handle "
        "high-volume card-present and card-not-present transactions for e-commerce and "
        "point-of-sale environments. The system interfaces with acquiring banks, card "
        "schemes (Visa, Mastercard, AMEX), and merchant back-office platforms to authorise, "
        "capture, and settle payment transactions in real time."
    )
    _para(doc,
        "The module supports full transaction lifecycle management including authorisation, "
        "capture, refunds, chargebacks, and daily settlement batch processing. It exposes "
        "a RESTful API consumed by merchant front-ends and integrates with the PayFlow "
        "Reconciliation Engine to ensure financial accuracy across all ledger entries. "
        "Availability SLA is 99.95 % with a maximum authorisation latency of 800 ms at P99."
    )

    # ── Functional Requirements ───────────────────────────────────────────────
    _heading(doc, "Functional Requirements", 1)

    _heading(doc, "Transaction Processing", 2)
    for req_id, text in [
        ("FR-001", "The system SHALL authorise card transactions within 800 ms under normal load."),
        ("FR-002", "The system SHALL support Visa, Mastercard, AMEX, and Maestro card schemes."),
        ("FR-003", "The system SHALL capture authorised transactions within the capture window "
                   "(default 7 days, configurable per merchant)."),
        ("FR-004", "The system SHALL reject duplicate transactions identified by idempotency key "
                   "within a 24-hour sliding window."),
        ("FR-005", "The system SHALL support 3-D Secure 2.x (3DS2) authentication flows for "
                   "card-not-present transactions."),
    ]:
        _para(doc, f"{req_id}: {text}")

    _heading(doc, "Refund Management", 2)
    for req_id, text in [
        ("FR-006", "The system SHALL process full and partial refunds against captured transactions."),
        ("FR-007", "Refunds SHALL be linked to the original transaction reference for audit trail."),
        ("FR-008", "The system SHALL prevent refunds exceeding the original transaction amount."),
        ("FR-009", "Refund status updates SHALL be propagated to the merchant webhook within 30 s."),
    ]:
        _para(doc, f"{req_id}: {text}")

    _heading(doc, "Reporting & Settlement", 2)
    for req_id, text in [
        ("FR-010", "The system SHALL generate daily settlement files in ISO 8583 and CSV formats."),
        ("FR-011", "Merchants SHALL be able to query transaction history with date-range and "
                   "status filters via the Merchant Portal API."),
        ("FR-012", "The Reconciliation Engine SHALL flag mismatches between captured amounts "
                   "and settled amounts and raise an alert to the Operations team."),
    ]:
        _para(doc, f"{req_id}: {text}")

    # ── Non-Functional Requirements ───────────────────────────────────────────
    _heading(doc, "Non-Functional Requirements", 1)
    _para(doc,
        "Performance: The system must sustain 2 000 transactions per second (TPS) peak load "
        "with P99 authorisation latency ≤ 800 ms. Horizontal scaling via Kubernetes HPA "
        "must activate within 60 s of load threshold being exceeded."
    )
    _para(doc,
        "Security: All data in transit must be encrypted with TLS 1.2+. PANs must be "
        "tokenised at the point of entry and never stored in plain text. The system must "
        "comply with PCI-DSS Level 1 requirements and pass quarterly ASV scans."
    )
    _para(doc,
        "Availability: The system must maintain 99.95 % uptime measured monthly. Planned "
        "maintenance windows must not exceed 4 hours per quarter and must be scheduled "
        "between 02:00–06:00 UTC."
    )

    # ── Glossary ──────────────────────────────────────────────────────────────
    _heading(doc, "Glossary", 1)
    _table(doc,
        headers=["Term", "Definition", "Module"],
        rows=[
            ["Transaction",
             "A financial exchange between a cardholder and merchant resulting in an authorisation, capture, or void.",
             "Core"],
            ["Settlement",
             "The process by which an acquirer transfers funds to the merchant's bank account after capture.",
             "Settlement"],
            ["Chargeback",
             "A forced reversal of a transaction initiated by the cardholder's issuing bank.",
             "Disputes"],
            ["Acquirer",
             "The bank or financial institution that processes card payments on behalf of the merchant.",
             "Core"],
            ["Issuer",
             "The bank or institution that issued the payment card to the cardholder.",
             "Core"],
            ["Merchant ID",
             "A unique identifier assigned to each merchant by the acquirer for routing and reporting.",
             "Onboarding"],
            ["Authorization",
             "A hold placed on cardholder funds to confirm sufficient balance before capture.",
             "Core"],
            ["Reconciliation",
             "The process of matching transaction records between PayFlow, the acquirer, and the merchant ledger.",
             "Settlement"],
        ],
    )

    # ── Domain Actors ─────────────────────────────────────────────────────────
    _heading(doc, "Domain Actors", 1)
    _table(doc,
        headers=["Actor", "Role", "Responsibilities"],
        rows=[
            ["Merchant",
             "Service provider",
             "Initiates transaction requests; receives settlement funds; manages refunds via Merchant Portal."],
            ["Cardholder",
             "End user",
             "Provides payment credentials; initiates purchases and refund requests."],
            ["Payment Gateway",
             "Intermediary system",
             "Routes authorisation requests to card schemes; enforces fraud rules; manages capture and settlement."],
            ["Bank (Acquirer / Issuer)",
             "Financial institution",
             "Authorises or declines transactions; settles funds; handles chargebacks."],
            ["QA Engineer",
             "Quality assurance",
             "Designs and executes test cases; validates FR coverage; reports defects; signs off releases."],
        ],
    )

    doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# 2. test_plan_payment.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_test_plan(path: Path):
    doc = docx.Document()
    doc.core_properties.title = "PayFlow Test Plan"
    doc.core_properties.author = "QA Team"

    doc.add_heading("PayFlow Payment Module — Test Plan v1.0", 0)

    # ── Test Scope ────────────────────────────────────────────────────────────
    _heading(doc, "Test Scope", 1)
    _para(doc,
        "In scope: Transaction authorisation (FR-001–FR-005), refund processing "
        "(FR-006–FR-009), settlement batch generation (FR-010), merchant reporting API "
        "(FR-011), and reconciliation mismatch alerting (FR-012). All environments "
        "(DEV, SIT, UAT) are included. 3DS2 authentication flows require the card "
        "scheme test harness provided by Visa and Mastercard."
    )
    _para(doc,
        "Out of scope: PCI-DSS compliance audit (handled by a separate Security team), "
        "performance benchmarking under production traffic (covered by the dedicated "
        "load-testing workstream), and any merchant-side front-end UI beyond API contract "
        "validation."
    )

    # ── Test Approach ─────────────────────────────────────────────────────────
    _heading(doc, "Test Approach", 1)

    _heading(doc, "Smoke Tests", 2)
    for item in [
        "Happy-path authorisation with a Visa test card — assert HTTP 200 and authorisation code returned.",
        "Single-item refund against a captured transaction — assert refund status APPROVED within 30 s.",
        "Merchant Portal: query last 7 days of transactions — assert response contains correct merchant_id filter.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _heading(doc, "Regression Tests", 2)
    for item in [
        "Full authorisation → capture → refund lifecycle for Visa, Mastercard, and AMEX cards.",
        "Duplicate transaction rejection: submit same idempotency key twice within 24 h — assert second call returns HTTP 409.",
        "Partial refund: refund 50 % of captured amount — assert refund amount and remaining balance are correct.",
        "Settlement batch generation at end-of-day: assert file contains all captured transactions and totals match ledger.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _heading(doc, "Integration Tests", 2)
    for item in [
        "End-to-end 3DS2 flow with Visa test harness: challenge flow and frictionless flow — assert both result in authorised transaction.",
        "Reconciliation Engine receives settlement file and raises alert on injected mismatch — assert Operations webhook fires within 60 s.",
        "Merchant webhook delivery on refund status change — assert payload schema and delivery retry on initial 500 response.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ── Test Environments ─────────────────────────────────────────────────────
    _heading(doc, "Test Environments", 1)
    _table(doc,
        headers=["Environment", "URL", "Purpose", "Owner"],
        rows=[
            ["DEV",          "https://payflow-dev.internal",  "Developer unit testing and integration smoke tests", "Engineering"],
            ["SIT",          "https://payflow-sit.internal",  "System integration testing with real card scheme simulators", "QA"],
            ["UAT",          "https://payflow-uat.internal",  "Business acceptance testing with merchant stakeholders", "Product"],
            ["PROD-mirror",  "https://payflow-mirror.prod",   "Production-like load and regression tests pre-release", "QA + Ops"],
        ],
    )

    # ── Entry & Exit Criteria ─────────────────────────────────────────────────
    _heading(doc, "Entry & Exit Criteria", 1)
    _para(doc,
        "Entry criteria: All FR-001–FR-012 requirements reviewed and signed off by Product; "
        "SIT environment deployed and smoke-tested green; card scheme test credentials "
        "provisioned; test data (merchant IDs, test PANs) loaded into SIT database."
    )
    _para(doc,
        "Exit criteria: 100 % of smoke tests passed; ≥ 95 % of regression test cases "
        "passed; zero open Critical or Major defects; UAT sign-off obtained from the "
        "Product Owner; release notes reviewed by QA Lead."
    )

    # ── Risk Register ─────────────────────────────────────────────────────────
    _heading(doc, "Risk Register", 1)
    _table(doc,
        headers=["Risk", "Probability", "Impact", "Mitigation"],
        rows=[
            ["Card scheme test harness unavailable during SIT window",
             "Medium", "High",
             "Book harness slots 2 weeks in advance; maintain fallback mock service."],
            ["3DS2 frictionless flow not testable in SIT (scheme restriction)",
             "High", "Medium",
             "Use Visa/Mastercard certified emulator; document limitation in test report."],
            ["Settlement file schema change from acquirer with short notice",
             "Low", "High",
             "Pin acquirer SDK version; add contract test in CI pipeline."],
            ["Test data PAN tokens expiring mid-regression causing false failures",
             "Medium", "Medium",
             "Automate token refresh in test setup hook; alert QA on token expiry."],
        ],
    )

    doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# 3. qa_process.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_qa_process(path: Path):
    doc = docx.Document()
    doc.core_properties.title = "QA Process & Procedures"
    doc.core_properties.author = "QA Lead"

    doc.add_heading("QA Process & Procedures — PayFlow Programme", 0)

    # ── QA Workflow Overview ──────────────────────────────────────────────────
    _heading(doc, "QA Workflow Overview", 1)
    _para(doc,
        "The PayFlow programme follows a V-model testing approach aligned to the "
        "ISTQB Foundation Level framework. Requirements are reviewed at inception, "
        "and corresponding test conditions are identified at each stage of the V. "
        "Unit tests are the responsibility of the development team; QA owns integration, "
        "system, and acceptance test levels. All test artefacts are stored in Jira/Xray "
        "and linked to the relevant Jira user stories. Test execution results are "
        "published to the QA dashboard after each test cycle."
    )

    # ── Defect Management ─────────────────────────────────────────────────────
    _heading(doc, "Defect Management", 1)

    _heading(doc, "Defect Lifecycle", 2)
    _para(doc,
        "All defects are raised in Jira with a mandatory template: Summary, Steps to "
        "Reproduce, Expected Result, Actual Result, Severity, Affected Environment, "
        "and Attachments (screenshot or HAR file). Defects are triaged at the daily "
        "stand-up and assigned to the responsible developer."
    )
    _table(doc,
        headers=["Status", "Description", "Owner"],
        rows=[
            ["New",         "Defect raised but not yet reviewed by the development team.", "QA Engineer"],
            ["Open",        "Defect acknowledged and accepted for fixing in current sprint.", "Dev Lead"],
            ["In Progress",  "Developer is actively working on the fix.", "Developer"],
            ["Fixed",       "Fix deployed to the target environment, awaiting QA verification.", "Developer"],
            ["Verified",    "QA has confirmed the fix resolves the defect on the target environment.", "QA Engineer"],
            ["Closed",      "Defect resolved and verified; closed after sprint retrospective.", "QA Lead"],
            ["Rejected",    "Defect raised in error or classified as expected behaviour.", "Dev Lead"],
        ],
    )

    _heading(doc, "Severity Levels", 2)
    _table(doc,
        headers=["Level", "Name", "Definition", "Response Time"],
        rows=[
            ["1", "Critical",
             "System is down or a core payment flow is completely broken. No workaround available.",
             "Fix within 4 hours; hotfix released same business day."],
            ["2", "Major",
             "Significant feature is impaired; workaround exists but is not acceptable long-term.",
             "Fix within 1 business day; included in next planned release."],
            ["3", "Minor",
             "Feature partially impaired with an acceptable workaround; low user impact.",
             "Fix within current sprint or next planned release."],
            ["4", "Trivial",
             "Cosmetic issue or minor UX inconsistency; no functional impact.",
             "Added to backlog; fixed at team's discretion."],
        ],
    )

    # ── Test Case Management ──────────────────────────────────────────────────
    _heading(doc, "Test Case Management", 1)

    _heading(doc, "Tagging Convention", 2)
    _para(doc,
        "Every test case in Xray must have at least one tag from the following "
        "taxonomy. Multiple tags are permitted."
    )
    for tag, desc in [
        ("smoke",        "Minimal set verifying the system starts and core flows respond."),
        ("regression",   "Full test suite executed on every release candidate."),
        ("integration",  "Tests spanning multiple services or external dependencies."),
        ("performance",  "Load and stress tests; run in PROD-mirror only."),
        ("security",     "Tests targeting OWASP Top-10 and PCI-DSS controls."),
        ("happy-path",   "The primary, error-free user journey for a feature."),
        ("negative",     "Input validation and error-handling scenarios."),
    ]:
        doc.add_paragraph(f"[{tag}] — {desc}", style="List Bullet")

    _heading(doc, "Review Process", 2)
    for step, text in [
        ("Step 1 — Peer Review",
         "Author submits test cases in Xray for peer review. Reviewer checks coverage "
         "against acceptance criteria, edge cases, and correct tag assignment."),
        ("Step 2 — QA Lead Approval",
         "QA Lead reviews high-severity and integration test cases. Approves or requests "
         "changes. Approval is mandatory before test cases can be executed in SIT."),
        ("Step 3 — Traceability Check",
         "QA Lead verifies every functional requirement (FR-001–FR-012) is covered by "
         "at least one test case. Gaps are raised as test-coverage defects."),
    ]:
        _para(doc, f"{step}: {text}")

    # ── Roles & Responsibilities ──────────────────────────────────────────────
    _heading(doc, "Roles & Responsibilities", 1)
    _table(doc,
        headers=["Role", "Responsibilities", "Tools"],
        rows=[
            ["QA Lead",
             "Owns the test strategy; approves test cases; chairs defect triage; "
             "signs off release readiness; reports to Programme Manager.",
             "Jira, Xray, Confluence, QA Dashboard"],
            ["QA Engineer",
             "Writes and executes test cases; raises defects; performs regression "
             "and integration testing; maintains test data.",
             "Jira, Xray, Postman, Selenium, k6"],
            ["Dev Lead",
             "Triages defects; ensures unit test coverage ≥ 80 %; reviews QA "
             "blockers and prioritises fixes.",
             "Jira, GitHub, SonarQube"],
            ["Product Owner",
             "Provides acceptance criteria; participates in UAT; signs off "
             "exit criteria before release.",
             "Jira, Confluence, Merchant Portal"],
        ],
    )

    doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# Verification
# ─────────────────────────────────────────────────────────────────────────────

async def verify(path: Path):
    # Add backend/ to sys.path so app imports work when run directly
    backend_dir = Path(__file__).resolve().parent.parent.parent
    if str(backend_dir) not in sys.path:
        sys.path.insert(0, str(backend_dir))

    from app.parsers.document_parser import DocumentParser
    result = await DocumentParser().parse(str(path))
    size_kb = path.stat().st_size / 1024
    return {
        "file":        path.name,
        "size_kb":     round(size_kb, 1),
        "text_len":    len(result["text"]),
        "headings":    len(result["headings"]),
        "tables":      len(result["tables"]),
        "text_ok":     len(result["text"]) > 0,
        "tables_ok":   len(result["tables"]) > 0,
    }


async def main():
    files = {
        "srs_payment_module.docx": build_srs,
        "test_plan_payment.docx":  build_test_plan,
        "qa_process.docx":         build_qa_process,
    }

    print("Generating synthetic docs …")
    for name, builder in files.items():
        p = OUT_DIR / name
        builder(p)
        print(f"  ✓ {name}")

    print("\nVerifying with DocumentParser …\n")
    print(f"{'File':<30} {'Size':>8} {'TextLen':>8} {'Headings':>9} {'Tables':>7} {'OK?':>5}")
    print("─" * 75)
    all_ok = True
    for name in files:
        v = await verify(OUT_DIR / name)
        ok = "✓" if (v["text_ok"] and v["tables_ok"]) else "✗"
        if not (v["text_ok"] and v["tables_ok"]):
            all_ok = False
        print(
            f"{v['file']:<30} {v['size_kb']:>7.1f}k "
            f"{v['text_len']:>8} {v['headings']:>9} {v['tables']:>7}   {ok}"
        )

    print()
    if all_ok:
        print("All files generated and verified successfully.")
    else:
        print("WARNING: one or more files failed verification.")
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
